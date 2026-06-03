"""
Calibrated Candidate - Document Extraction Pipeline (Skill 02)
Normalizes uploaded documents to clean Markdown before they enter conversation.

Method chain:
  1. markitdown        (PDF text-layer primary)
  2. docx_native       (DOCX via python-docx, including tables)
  3. ocr_rasterized    (scanned PDF -> pdfplumber rasterize -> pytesseract)
  4. ocr_direct        (image uploads -> pytesseract)
  5. user_paste        (all other failures -> caller prompts user)

Design rule: never raise a traceback to the user. Every failure returns a
structured result with a plain-language message and routes to the next method.

DOCX is read with python-docx directly rather than through markitdown's
optional [docx] extra, so DOCX support does not depend on how markitdown
was installed in any given container.
"""
import os, datetime

MIN_CHARS = 100  # below this, treat a "successful" PDF parse as scanned/empty

def _now():
    return datetime.datetime.now(datetime.timezone.utc).isoformat()

def _result(markdown, method, ok, message=""):
    return {
        "markdown": markdown or "",
        "extraction_method": method,
        "extracted_at": _now() if ok else "",
        "ok": ok,
        "message": message,
        "char_count": len(markdown or ""),
    }

# ---- Runtime guard: detect missing libraries up front, plain language ----
def check_environment():
    missing = []
    for mod, label in [("markitdown","markitdown"),("pdfplumber","pdfplumber"),
                        ("docx","python-docx"),("PIL","Pillow"),
                        ("pytesseract","pytesseract")]:
        try:
            __import__(mod)
        except Exception:
            missing.append(label)
    tess_ok = True
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
    except Exception:
        tess_ok = False
    return {"missing_libraries": missing, "tesseract_available": tess_ok}

# ---- Method: DOCX via python-docx (paragraphs + tables) ----
def try_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if text:
            return _result(text, "docx_native", True)
        return _result("", "docx_native", False, "This Word file appears to be empty.")
    except Exception as e:
        return _result("", "docx_native", False,
                       f"Could not read this Word file ({type(e).__name__}).")

# ---- Method: markitdown (PDF primary) ----
def try_markitdown(path):
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        out = md.convert(path)
        text = (out.text_content or "").strip()
        if len(text) >= MIN_CHARS:
            return _result(text, "markitdown", True)
        return _result(text, "markitdown", False,
                       "Low text yield, will try OCR fallback.")
    except Exception as e:
        return _result("", "markitdown", False,
                       f"Primary reader could not open this file ({type(e).__name__}).")

# ---- Method: OCR a scanned PDF by rasterizing pages ----
def try_ocr_pdf(path):
    try:
        import pdfplumber, pytesseract
        pages_text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                im = page.to_image(resolution=200).original
                pages_text.append(pytesseract.image_to_string(im))
        text = "\n\n".join(pages_text).strip()
        if text:
            return _result(text, "ocr_rasterized", True,
                           "Used OCR; accuracy may vary on low-quality scans.")
        return _result("", "ocr_rasterized", False, "OCR produced no text.")
    except Exception as e:
        return _result("", "ocr_rasterized", False,
                       f"OCR on this PDF failed ({type(e).__name__}).")

# ---- Method: OCR a direct image ----
def try_ocr_image(path):
    try:
        from PIL import Image
        import pytesseract
        im = Image.open(path)
        raw = pytesseract.image_to_string(im)
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        text = "\n".join(lines).strip()
        if text:
            return _result(text, "ocr_direct", True,
                           "Used OCR on an image; accuracy may vary.")
        return _result("", "ocr_direct", False, "OCR produced no text.")
    except Exception as e:
        return _result("", "ocr_direct", False,
                       f"OCR on this image failed ({type(e).__name__}).")

IMAGE_EXTS = {".png",".jpg",".jpeg",".tif",".tiff",".bmp",".gif"}

def _looks_like_zip(path):
    # DOCX files are ZIP archives; first bytes are 'PK'. Used to catch a real
    # DOCX that arrived with a wrong extension before we mishandle it as PDF.
    try:
        with open(path, "rb") as f:
            return f.read(2) == b"PK"
    except Exception:
        return False

def extract(path):
    """Top-level entry. Returns a structured result, never raises to caller."""
    if not os.path.exists(path):
        return _result("", "user_paste", False, "File not found at the given path.")
    if os.path.getsize(path) == 0:
        return _result("", "user_paste", False,
                       "This file is empty. Please paste your text directly.")
    ext = os.path.splitext(path)[1].lower()

    # Image: OCR direct
    if ext in IMAGE_EXTS:
        r = try_ocr_image(path)
        if r["ok"]: return r
        return _result("", "user_paste", False,
                       "Could not read this image. Please paste your text directly.")

    # DOCX: markitdown first (richer table structure when its docx support is
    # present), then python-docx as a reliable fallback when it is not.
    if ext == ".docx":
        rm = try_markitdown(path)
        if rm["ok"]: return rm
        r = try_docx(path)
        if r["ok"]: return r
        return _result("", "user_paste", False,
                       "Could not read this Word file. Please paste your text directly.")

    # PDF and everything else: markitdown first
    r = try_markitdown(path)
    if r["ok"]:
        return r

    # If a non-.docx file is actually a ZIP/DOCX in disguise, try python-docx
    if ext != ".docx" and _looks_like_zip(path):
        rd = try_docx(path)
        if rd["ok"]:
            return rd

    # markitdown weak or failed; if PDF, try OCR rasterize
    if ext == ".pdf":
        r2 = try_ocr_pdf(path)
        if r2["ok"]: return r2

    # everything failed -> paste fallback
    return _result("", "user_paste", False,
                   "Could not read this file automatically. Please paste your text directly.")

if __name__ == "__main__":
    import sys, json
    print(json.dumps(extract(sys.argv[1]), indent=2))
